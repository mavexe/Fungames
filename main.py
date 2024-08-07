import tkinter as tk
from tkinter import filedialog, messagebox
import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

class ExcelToWordApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Excel to Word")
        self.excel_file = None
        self.template_file = None

        self.create_widgets()

    def create_widgets(self):
        # Label and button to select Excel file
        self.excel_label = tk.Label(self.root, text="Select Excel File:")
        self.excel_label.grid(row=0, column=0, padx=10, pady=10)
        
        self.excel_button = tk.Button(self.root, text="Browse", command=self.select_excel_file)
        self.excel_button.grid(row=0, column=1, padx=10, pady=10)
        
        # Label and button to select Word template
        self.template_label = tk.Label(self.root, text="Select Word Template:")
        self.template_label.grid(row=1, column=0, padx=10, pady=10)
        
        self.template_button = tk.Button(self.root, text="Browse", command=self.select_template_file)
        self.template_button.grid(row=1, column=1, padx=10, pady=10)

        # Button to edit Word template
        self.edit_button = tk.Button(self.root, text="Edit Template", command=self.edit_template)
        self.edit_button.grid(row=2, column=0, columnspan=2, pady=10)

        # Button to generate Word document
        self.generate_button = tk.Button(self.root, text="Generate Word Document", command=self.generate_word_document)
        self.generate_button.grid(row=3, column=0, columnspan=2, pady=10)

    def select_excel_file(self):
        self.excel_file = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx *.xls")])
        if self.excel_file:
            messagebox.showinfo("Selected Excel File", self.excel_file)

    def select_template_file(self):
        self.template_file = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if self.template_file:
            messagebox.showinfo("Selected Word Template", self.template_file)

    def read_excel(self, file_path):
        wb = openpyxl.load_workbook(file_path)
        sheet = wb.active
        data = {}
        for row in sheet.iter_rows(values_only=True):
            if not data:
                headers = row
                for header in headers:
                    data[header] = []
            else:
                for header, value in zip(headers, row):
                    data[header].append(value)
        return data

    def fill_word_template(self, template_path, output_path, replacements):
        doc = Document(template_path)
        
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    inline = paragraph.runs
                    for item in inline:
                        if key in item.text:
                            item.text = item.text.replace(key, str(value))
        
        # Save the modified document
        doc.save(output_path)

    def generate_word_document(self):
        if not self.excel_file or not self.template_file:
            messagebox.showwarning("Input Error", "Please select both Excel file and Word template.")
            return
        
        try:
            data = self.read_excel(self.excel_file)
        except Exception as e:
            messagebox.showerror("Error", f"Error reading Excel file: {str(e)}")
            return
        
        try:
            output_file = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if output_file:
                for row_index in range(len(next(iter(data.values())))):
                    replacements = {f'{{{key}}}': data[key][row_index] for key in data}
                    self.fill_word_template(self.template_file, output_file, replacements)
                messagebox.showinfo("Success", "Word document has been generated successfully.")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def edit_template(self):
        if not self.template_file:
            messagebox.showwarning("Input Error", "Please select a Word template to edit.")
            return

        edit_window = tk.Toplevel(self.root)
        edit_window.title("Edit Template")
        text = tk.Text(edit_window, wrap=tk.WORD)
        text.pack(expand=True, fill=tk.BOTH)

        # Load the content of the Word template into the text widget
        try:
            doc = Document(self.template_file)
            content = []
            for paragraph in doc.paragraphs:
                content.append(paragraph.text)
            text.insert(tk.END, "\n".join(content))
        except Exception as e:
            messagebox.showerror("Error", f"Error loading Word template: {str(e)}")
            edit_window.destroy()
            return

        def save_template():
            edited_content = text.get("1.0", tk.END).splitlines()
            try:
                doc = Document()
                for line in edited_content:
                    doc.add_paragraph(line)
                doc.save(self.template_file)
                messagebox.showinfo("Success", "Template has been updated successfully.")
                edit_window.destroy()
            except Exception as e:
                messagebox.showerror("Error", f"Error saving Word template: {str(e)}")

        save_button = tk.Button(edit_window, text="Save", command=save_template)
        save_button.pack(pady=10)

if __name__ == "__main__":
    root = tk.Tk()
    app = ExcelToWordApp(root)
    root.mainloop()
